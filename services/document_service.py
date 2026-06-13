
# ═══ 必须在所有 Paddle 相关 import 之前禁用 oneDNN ═══
import os
os.environ["FLAGS_use_onednn"] = "0"
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
# ════════════════════════════════════════════════════════

from sqlalchemy.exc import SQLAlchemyError
from models import Document, DocumentChunk
from pathlib import Path
import uuid
import re
import tempfile
from logging_config import get_logger

logger = get_logger(__name__)

# PDF 解析
from pypdf import PdfReader

# Word 解析
from docx import Document as DocxDocument

# OCR 相关（懒加载，避免启动时就加载大模型）
_fitz = None
_paddleocr_engine = None
_ocr_degraded = False  # 标记 OCR 是否已降级

def _get_fitz():
    """懒加载 PyMuPDF (fitz)，避免启动时加载"""
    global _fitz
    if _fitz is None:
        import fitz
        _fitz = fitz
    return _fitz

def _get_paddleocr():
    """
    懒加载 PaddleOCR 引擎（线程安全）。
    oneDNN 已在模块级别通过环境变量 FLAGS_use_onednn=0 禁用。
    KMP_DUPLICATE_LIB_OK=TRUE 防止 Windows 下 OpenMP 冲突。
    """
    global _paddleocr_engine, _ocr_degraded
    import threading

    # 双重检查锁定，防止 4 个 OCR 工作线程同时初始化导致 segfault
    if _paddleocr_engine is None and not _ocr_degraded:
        # 使用模块级锁
        lock_attr = '_paddleocr_init_lock'
        if not hasattr(_get_paddleocr, lock_attr):
            setattr(_get_paddleocr, lock_attr, threading.Lock())
        lock = getattr(_get_paddleocr, lock_attr)

        with lock:
            if _paddleocr_engine is None:  # 双重检查
                try:
                    from paddleocr import PaddleOCR
                    _paddleocr_engine = PaddleOCR(lang='ch', use_angle_cls=True)
                    logger.info("PaddleOCR 引擎初始化完成")
                except Exception as e:
                    logger.error(f"PaddleOCR 初始化失败: {e}，将禁用 OCR 功能")
                    _ocr_degraded = True
                    _paddleocr_engine = None
    return _paddleocr_engine

# jieba 分词
import jieba
import jieba.posseg as pseg

# embedding 服务
import services.embedding_service
import traceback
from sqlalchemy import or_

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)


# ═══════════════════════════════════════════════════════════════
# OCR 工具函数
# ═══════════════════════════════════════════════════════════════

def _is_text_garbled(text: str) -> bool:
    """
    检测文本是否为乱码（增强版 v2）。
    判断标准：
    1. 文本为空 → True
    2. 有意义字符（中文+英文+数字）占比 < 40% → True（从 20% 提升）
    3. 中文字符占比极低（< 5%）且无英文单词 → True
    4. 纯数字/符号表格（数字占比 > 30%）→ False，不是乱码
    5. 连续无效字符序列 > 40 → True（段落级乱码检测）
    6. 伪 CJK 字符占比检测（GBK 双重编码等产生的高码点字符）→ True
    """
    if not text or text.strip() == "":
        return True

    text = text.strip()
    total = len(text)

    # 7. PDF 字形 ID 乱码检测：/G21 /G2A /G1F 等模式
    # 当 pypdf 无法解码实际文字时，会返回字体字形 ID
    glyph_patterns = len(re.findall(r'/G[0-9A-Fa-f]{1,3}', text))
    if glyph_patterns >= 5 and glyph_patterns * 4 >= total * 0.3:
        return True

    chinese = 0
    english_alpha = 0
    digit = 0
    whitespace = 0
    # 新增统计
    pseudo_cjk = 0       # 伪 CJK 字符（CJK 扩展 B+ 区，实际文本中极少出现）
    consecutive_invalid = 0  # 连续无效字符计数
    max_consecutive_invalid = 0

    # 合法的常见标点符号（中文+英文）
    valid_punctuation = set(
        '，。！？；：""''（）【】《》、…—\n\r\t　'
        ',.!?;:\'\"()[]{}<>/-_=+@#$%^&*|\\~`'
        '·•·∙※←→↑↓↔↵✓✗✘□■○●◇◆△▲☆★▷▶▸►'
        '①②③④⑤⑥⑦⑧⑨⑩ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ'
    )

    for ch in text:
        cp = ord(ch)
        # 中文字符范围（含 CJK 统一汉字、扩展 A-F、兼容汉字）
        if (0x4E00 <= cp <= 0x9FFF or      # CJK 统一汉字 (常用)
            0x3400 <= cp <= 0x4DBF):       # CJK 扩展 A
            chinese += 1
            consecutive_invalid = 0
        elif (0xF900 <= cp <= 0xFAFF):     # CJK 兼容汉字
            chinese += 1
            consecutive_invalid = 0
        elif (0x20000 <= cp <= 0x2A6DF):   # CJK 扩展 B (极少使用，可能是伪字符)
            pseudo_cjk += 1
            consecutive_invalid = 0
        elif ('a' <= ch.lower() <= 'z'):
            english_alpha += 1
            consecutive_invalid = 0
        elif ('0' <= ch <= '9'):
            digit += 1
            consecutive_invalid = 0
        elif ch in ' \n\r\t　':
            whitespace += 1
            # 空白不重置 consecutive_invalid，也不累加
            continue
        elif ch in valid_punctuation:
            # 标点不算乱码但也不累加连续计数
            consecutive_invalid = 0
        else:
            # 无法识别的字符 → 可能是乱码
            consecutive_invalid += 1
            if consecutive_invalid > max_consecutive_invalid:
                max_consecutive_invalid = consecutive_invalid

    meaningful = chinese + english_alpha + digit
    non_whitespace = total - whitespace

    if non_whitespace == 0:
        return True

    # 有意义字符占非空白字符的比例
    meaningful_ratio = meaningful / non_whitespace

    # 中文占比
    chinese_ratio = chinese / non_whitespace

    # 数字占比
    digit_ratio = digit / non_whitespace

    # 伪 CJK 占比
    pseudo_cjk_ratio = pseudo_cjk / non_whitespace

    # ── 判断条件 ──

    # 1. 纯数字/符号表格：数字占比 > 30% → 不是乱码（表格数据）
    #    但如果是 /G21/G22 类编码垃圾模式，即使数字多也是乱码
    slash_like_codes = len(re.findall(r'/[A-Z]\d{2,}', text))
    if digit_ratio > 0.30 and slash_like_codes <= 10:
        return False

    # 2. 有意义字符 < 40%（提升阈值，原来 20% 太宽松）→ 乱码
    if meaningful_ratio < 0.40:
        return True

    # 3. 既无中文也无英文 → 乱码（纯符号或损坏数据）
    if chinese == 0 and english_alpha == 0:
        return True

    # 4. 有意义字符 < 50% 且中文 < 5% → 乱码
    if meaningful_ratio < 0.50 and chinese_ratio < 0.05:
        return True

    # 5. 连续无效字符超过 40 个 → 文本中有大片乱码段
    if max_consecutive_invalid > 40:
        return True

    # 6. 伪 CJK 字符占比 > 15%（CJK 扩展 B 区字符大量出现通常是编码错误）
    if pseudo_cjk_ratio > 0.15:
        return True

    # 7. 重复垃圾模式检测：/G21/G22/G23 类伪编码
    #    提取所有连续英文数字串，检查是否几乎都是 "字母+数字" 模式
    alpha_num_tokens = re.findall(r'[A-Za-z]+\d+|\d+[A-Za-z]+', text)
    if len(alpha_num_tokens) > 20:
        # 如果有大量 "字母+数字" token 且中文极少 → 编码垃圾
        if chinese_ratio < 0.10:
            return True

    # 8. 可疑重复编码检测：/G21/G22... 模式 + 中文极少 → 垃圾
    slash_code_pattern = re.findall(r'/[A-Z]\d{2,}', text)
    if len(slash_code_pattern) > 10 and chinese_ratio < 0.15:
        return True

    # 9. 超低信息密度：中文 < 3% 且有大量单字母 token → 垃圾
    if chinese_ratio < 0.03 and english_alpha > 100:
        en_tokens = re.findall(r'[A-Za-z]+', text)
        if en_tokens:
            avg_len = sum(len(t) for t in en_tokens) / len(en_tokens)
            if avg_len < 2.5:
                return True

    return False


def _ocr_pdf_with_paddle(file_path: str) -> str:
    """
    使用 PaddleOCR + PyMuPDF 对 PDF 进行 OCR 识别（并行版 + 缓存）。
    - 并行 OCR 多页（ThreadPoolExecutor，4 线程）
    - 智能降 DPI：普通页 200 DPI，复杂页 300 DPI
    - OCR 结果写入 ocr_result_cache 表，相同页面复用
    """
    import hashlib
    import threading
    from concurrent.futures import ThreadPoolExecutor, as_completed

    fitz = _get_fitz()
    doc = fitz.open(file_path)
    total_pages = len(doc)
    texts = [None] * total_pages

    # 在主线程中预加载 PaddleOCR + 创建串行锁（PaddleOCR 不是线程安全的）
    _ocr_init_lock = threading.Lock()
    _ocr_call_lock = threading.Lock()
    with _ocr_init_lock:
        ocr_engine = _get_paddleocr()
    if ocr_engine is None:
        doc.close()
        return ""  # OCR 不可用，返回空

    # 尝试连接数据库做缓存
    db_session = None
    try:
        from database import SessionLocal
        from models import OcrResultCache
        db_session = SessionLocal()
    except Exception:
        pass  # 无数据库时跳过缓存

    def _ocr_single_page(page_index: int) -> tuple[int, str | None]:
        """OCR 单个页面（线程安全）"""
        page = doc[page_index]

        # 预检测页面复杂度：文本密集 → 需要更高 DPI
        # 简单启发式：页面尺寸大的可能是图纸/表格，用 300 DPI
        rect = page.rect
        page_area = rect.width * rect.height
        dpi = 300 if page_area > 600000 else 200  # A4约 595*842 = 500990

        pix = page.get_pixmap(dpi=dpi)

        # 计算图像 hash 用于缓存
        img_bytes = pix.tobytes("png")
        img_hash = hashlib.md5(img_bytes).hexdigest()

        # 检查缓存
        cached_text = None
        if db_session:
            try:
                cached = db_session.query(OcrResultCache).filter(
                    OcrResultCache.content_hash == img_hash
                ).first()
                if cached:
                    cached_text = cached.ocr_text
            except Exception:
                pass

        if cached_text:
            return (page_index, cached_text)

        # 执行 OCR（串行化，PaddleOCR 不线程安全）
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
            tmp.write(img_bytes)

        try:
            with _ocr_call_lock:
                result = ocr_engine.ocr(tmp_path)
            page_text_parts = []

            if result and result[0]:
                for line in result[0]:
                    text = line[1][0]
                    page_text_parts.append(text)

            page_text = " ".join(page_text_parts).strip()

            # 写入缓存
            if page_text and db_session:
                try:
                    db_session.add(OcrResultCache(
                        content_hash=img_hash,
                        ocr_text=page_text,
                        model_name="paddleocr",
                        page_count=1,
                    ))
                    db_session.commit()
                except Exception:
                    db_session.rollback()

            return (page_index, page_text if page_text else None)

        finally:
            Path(tmp_path).unlink(missing_ok=True)

    # 并行 OCR（4 线程）
    max_workers = min(4, total_pages)
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_ocr_single_page, i): i for i in range(total_pages)}
        for future in as_completed(futures):
            idx, text = future.result()
            texts[idx] = text

    doc.close()
    if db_session:
        db_session.close()

    # 按页码拼接
    result_parts = []
    for i, text in enumerate(texts):
        if text:
            result_parts.append(f"\n\n===== 第 {i + 1} 页 =====\n{text}")

    return "\n".join(result_parts)


def _ocr_single_image(image_path: str) -> str | None:
    """
    对单张图片执行 OCR，返回识别文本。
    如果 OCR 引擎不可用或识别失败，返回 None。
    """
    global _ocr_degraded

    if _ocr_degraded:
        return None

    ocr = _get_paddleocr()
    if ocr is None:
        return None

    try:
        result = ocr.ocr(image_path)
        if result and result[0]:
            parts = [line[1][0] for line in result[0]]
            return " ".join(parts).strip()
    except Exception as e:
        logger.warning(f"单图 OCR 失败: {e}")
        _ocr_degraded = True
    return None


def _extract_docx_images(doc) -> list[str]:
    """
    从 DOCX 文档中提取嵌入图片并 OCR。
    返回 OCR 识别文本列表，每个元素对应一张图片的识别结果。
    如果无法提取图片或 OCR 不可用，返回空列表。
    """
    global _ocr_degraded

    if _ocr_degraded:
        return []

    ocr = _get_paddleocr()
    if ocr is None:
        return []

    image_texts = []
    try:
        # 遍历 docx 内部的图片关系
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_part = rel.target_part
                image_bytes = image_part.blob
                # 保存到临时文件
                suffix = ".png"
                if hasattr(image_part, 'partname') and image_part.partname:
                    ext = Path(str(image_part.partname)).suffix.lower()
                    if ext in ('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff'):
                        suffix = ext
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                    tmp.write(image_bytes)
                    tmp_path = tmp.name
                try:
                    ocr_text = _ocr_single_image(tmp_path)
                    if ocr_text:
                        image_texts.append(ocr_text)
                finally:
                    Path(tmp_path).unlink(missing_ok=True)
    except Exception as e:
        logger.warning(f"提取嵌入图片失败: {e}")

    return image_texts


# ═══════════════════════════════════════════════════════════════
# 文档 CRUD 函数
# ═══════════════════════════════════════════════════════════════

def find_document(db, filename):
    existing_document = db.query(Document).filter(Document.filename == filename).first()
    return existing_document


def create_document(db, filename, standard_type, industry, tags, filepath: str | None = None):
    """创建文档记录。tags 支持 list[str] 或逗号分隔字符串。"""
    if isinstance(tags, str):
        tags = [t.strip() for t in tags.split(",") if t.strip()]
    try:
        new_document = Document(
            filename=filename,
            filepath=filepath,
            standard_type=standard_type,
            industry=industry,
            tags=tags,
        )
        db.add(new_document)
        db.commit()
        db.refresh(new_document)
    except SQLAlchemyError:
        db.rollback()
        return None
    return new_document


async def get_file(file):
    allowed_extensions = {".pdf", ".docx", ".txt"}
    original_filename = file.filename
    if original_filename is None:
        return None
    suffix = Path(original_filename).suffix.lower()
    if suffix not in allowed_extensions:
        return None
    file_bytes = await file.read()
    logger.info(f"上传文件读取字节数：{len(file_bytes)}")
    saved_name = uuid.uuid4().hex + suffix
    filepath = UPLOAD_DIR / saved_name
    filepath.write_bytes(file_bytes)
    return str(filepath)


def get_documents(db, filename, standard_type, industry, tag, page, page_size):
    query = db.query(Document)
    if filename is not None:
        query = query.filter(Document.filename == filename)
    if standard_type is not None:
        query = query.filter(Document.standard_type == standard_type)
    if industry is not None:
        query = query.filter(Document.industry == industry)
    if tag is not None:
        query = query.filter(Document.tags == tag)
    total_count = query.count()
    start = (page - 1) * page_size
    page_documents = query.offset(start).limit(page_size).all()
    return total_count, page_documents


def get_document_stats(db):
    standard_types = {}
    industries = {}
    tag_stats = {}
    documents = db.query(Document).all()
    for document in documents:
        standard_type = document.standard_type
        if standard_type is not None:
            standard_types[standard_type] = standard_types.get(standard_type, 0) + 1
        industry = document.industry
        if industry is not None:
            industries[industry] = industries.get(industry, 0) + 1
        tags = document.tags or []
        for tag in tags:
            tag_stats[tag] = tag_stats.get(tag, 0) + 1
    total = len(documents)
    return total, standard_types, industries, tag_stats


def search_documents(db, keyword, page, page_size):
    search_result = []
    documents = db.query(Document).all()
    keyword_lower = keyword.lower()
    for document in documents:
        match_fields = []
        if keyword_lower in (document.filename or "").lower():
            match_fields.append("filename")
        if keyword_lower in (document.standard_type or "").lower():
            match_fields.append("standard_type")
        if keyword_lower in (document.industry or "").lower():
            match_fields.append("industry")
        if any(keyword_lower in tag.lower() for tag in document.tags or []):
            match_fields.append("tags")
        if match_fields:
            match_document = {
                "id": document.id,
                "filename": document.filename,
                "standard_type": document.standard_type,
                "industry": document.industry,
                "tags": document.tags,
                "match_fields": match_fields,
            }
            search_result.append(match_document)
    start = (page - 1) * page_size
    end = start + page_size
    page_documents = search_result[start:end]
    has_more = end < len(search_result)
    total_count = len(page_documents)
    return keyword, total_count, has_more, page_documents


def get_document_by_id(db, document_id: int):
    document = db.query(Document).filter(Document.id == document_id).first()
    return document


def patch_document(db, filename, standard_type, industry, tags, document):
    try:
        if filename is not None:
            document.filename = filename
        if standard_type is not None:
            document.standard_type = standard_type
        if industry is not None:
            document.industry = industry
        if tags is not None:
            document.tags = tags
        db.commit()
        db.refresh(document)
    except SQLAlchemyError:
        db.rollback()
        return None
    return document


def delete_document(db, d_document):
    try:
        # 级联删除关联的 chunks
        db.query(DocumentChunk).filter(DocumentChunk.document_id == d_document.id).delete()
        # 级联删除关联的分析记录
        from models import DocumentAnalysis
        db.query(DocumentAnalysis).filter(DocumentAnalysis.document_id == d_document.id).delete()
        # 删除文档本身
        db.delete(d_document)
        db.commit()
    except SQLAlchemyError:
        db.rollback()
        return False
    return True


# ═══════════════════════════════════════════════════════════════
# 文档解析 (PDF / DOCX / TXT)
# ═══════════════════════════════════════════════════════════════

def _extract_docx_tables(doc: DocxDocument) -> list[str]:
    """
    从 python-docx 的 Document 对象中提取所有表格内容。
    返回格式化字符串列表，每个表格一个字符串。
    """
    table_texts = []

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                # 获取单元格文本，去除内部换行保持整洁
                cell_text = cell.text.strip().replace("\n", " ")
                cells.append(cell_text)
            rows.append(" | ".join(cells))

        if rows:
            table_text = (
                f"\n\n===== [表格 {table_index}] =====\n"
                + "\n".join(rows)
            )
            table_texts.append(table_text)

    return table_texts


def _extract_docx_textboxes(doc: DocxDocument) -> list[str]:
    """
    尝试从 docx XML body 中提取文本框 (w:txbxContent) 内的文字。
    如果提取不到，返回空列表。
    """
    textbox_texts = []
    try:
        body = doc.element.body
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'v': 'urn:schemas-microsoft-com:vml',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        }

        # 查找所有 txbxContent 元素
        txbx_elements = body.findall('.//w:txbxContent', nsmap)
        for elem in txbx_elements:
            # 提取 <w:t> 标签中的文本
            t_elements = elem.findall('.//w:t', nsmap)
            text = "".join(t.text or "" for t in t_elements).strip()
            if text:
                textbox_texts.append(text)
    except Exception:
        pass  # XML 解析失败就跳过文本框提取

    return textbox_texts


def parse_document(filepath):
    """
    解析上传的文档（.txt / .pdf / .docx），返回纯文本字符串。
    - PDF: 先尝试提取文本层，若文本为空或乱码率过高则自动降级到 PaddleOCR
    - DOCX: 提取段落 + 表格 + 文本框 + 嵌入图片 OCR
    - TXT: 直接读取
    """
    global _ocr_degraded
    file_path = Path(filepath)
    suffix = file_path.suffix.lower()

    # ── TXT ───────────────────────────────────────────────
    if suffix == ".txt":
        text = file_path.read_text(encoding="utf-8")
        logger.info(f"解析文本长度：{len(text)}")
        return text

    # ── PDF ───────────────────────────────────────────────
    if suffix == ".pdf":
        reader = PdfReader(file_path)
        texts = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                texts.append(f"\n\n===== 第 {index} 页 =====\n{page_text}")

        text = "\n".join(texts)

        # 检测是否乱码或空文本
        if _is_text_garbled(text):
            logger.warning(f"PDF 文本层乱码/为空，降级到 PaddleOCR: {file_path}")

            # 检查 OCR 引擎是否可用
            if _ocr_degraded:
                logger.warning("OCR 引擎已降级，返回原始文本（可能含乱码），请手动检查")
                # 返回原始文本但带上警告标记
                warning = (
                    "【警告】该 PDF 文本层可能包含乱码，且 OCR 引擎不可用。"
                    "以下内容仅供参考，建议人工核对原文。\n\n"
                )
                return warning + text

            try:
                ocr_text = _ocr_pdf_with_paddle(str(file_path))
                if ocr_text.strip():
                    logger.info(f"OCR 成功，识别文本长度: {len(ocr_text)}")
                    return ocr_text
                else:
                    logger.warning("OCR 识别结果也为空")
            except Exception as e:
                logger.error(f"PaddleOCR 识别失败: {e}")
                traceback.print_exc()
                # 标记 OCR 降级
                _ocr_degraded = True

            # 文本乱码且 OCR 失败 → 返回原始文本 + 警告，不抛异常阻塞流程
            warning = (
                "【警告】该 PDF 文本层包含乱码且 OCR 识别失败。"
                "以下内容可能不完整，建议核实原文。\n\n"
            )
            return warning + text

        return text

    # ── DOCX ───────────────────────────────────────────────
    if suffix == ".docx":
        doc = DocxDocument(str(file_path))
        text_parts = []

        # 1. 提取段落文本
        paragraph_texts = []
        for para in doc.paragraphs:
            para_text = para.text
            if para_text and para_text.strip():
                paragraph_texts.append(para_text.strip())
        if paragraph_texts:
            text_parts.append("\n\n".join(paragraph_texts))

        # 2. 提取表格文本
        table_texts = _extract_docx_tables(doc)
        if table_texts:
            text_parts.extend(table_texts)

        # 3. 提取文本框内容
        textbox_texts = _extract_docx_textboxes(doc)
        if textbox_texts:
            for tb_text in textbox_texts:
                text_parts.append(f"\n\n===== [文本框] =====\n{tb_text}")

        # 4. 提取嵌入图片并 OCR（兜底：扫描件 DOCX）
        try:
            docx_image_texts = _extract_docx_images(doc)
            if docx_image_texts:
                for idx, img_text in enumerate(docx_image_texts, start=1):
                    text_parts.append(f"\n\n===== [文档图片 {idx} OCR] =====\n{img_text}")
                logger.info(f"从 {len(docx_image_texts)} 张嵌入图片中 OCR 提取了文本")
        except Exception as e:
            logger.warning(f"嵌入图片 OCR 失败（非致命）: {e}")

        text = "\n".join(text_parts)

        # 如果全部为空，尝试仅 OCR 图片（完全扫描型 DOCX）
        if text.strip() == "":
            raise ValueError(
                "文档解析结果为空。该文档可能为纯扫描图片格式，"
                "或包含复杂排版（文本框、图片文字）暂不支持。"
                "建议转换为文本 PDF 或启用 OCR 后重新上传。"
            )

        return text

    return None


# ═══════════════════════════════════════════════════════════════
# 文档结构解析引擎 v2
# ═══════════════════════════════════════════════════════════════

# ── 正则可复用 ──
_PAGE_PATTERN = re.compile(r'=====\s*第\s*(\d+)\s*页\s*=====')
_TABLE_PATTERN = re.compile(r'=====\s*\[表格\s*(\d+)\]\s*=====')
_TEXTBOX_PATTERN = re.compile(r'=====\s*\[文本框\]\s*=====')

# 标准编号（GB/T 46350-2025、NY/T 898、ISO 9001:2015 等）
_STD_NUM_RE = re.compile(
    r'(?:GB[/TZ]?|GJB|NY[/T]?|HJ[/T]?|SN[/T]?|YY[/T]?|QB[/T]?|'
    r'DB\d+[/T]?|T/|ISO[/T]?|IEC[/T]?|ASTM[/T]?|'
    r'CJJ|JGJ|CECS|GA|SY|DL|WS|MH|JT|LS|LY|SC|SB|'
    r'Q/)\s*[/T]*\s*[\d\.]+(?:[—\-]\d+)?(?::\d+)?',
    re.IGNORECASE
)

# 条款编号（4、4.1、4.1.2、A.1）
_CLAUSE_LINE_RE = re.compile(
    r'^(\d+(?:\.\d+)*)\s+(.+)$'
)
_APPENDIX_HEADER_RE = re.compile(r'^(附录\s*[A-Za-z])\s*(.*)$')

# 正文各级标题匹配 (非条款编号的标题，如"总则"、"技术要求")
_CHAPTER_NO_RE = re.compile(r'^(\d+)\s+(.+)$')

# 表格/图在文本中的标记
_TABLE_TITLE_RE = re.compile(r'(?:^|\n)(表\s*\d+)[\s\.、:：—\-]*(.*?)(?:\n|$)')
_FIGURE_TITLE_RE = re.compile(r'(?:^|\n)(图\s*\d+)[\s\.、:：—\-]*(.*?)(?:\n|$)')

# 术语条目（3.1\n术语名称\n定义 或 3.1 术语名称\n定义）
_TERM_ENTRY_RE = re.compile(r'^(\d+\.\d+)\s+(.+)$', re.MULTILINE)


def _map_page_numbers(text: str) -> tuple[list[dict], dict[int, int]]:
    """
    通过页码标记切分文本，建立「字符位置→页码」映射。
    返回: (pages列表, char_to_page映射)
    """
    pages = []
    char_to_page = {}

    page_matches = list(_PAGE_PATTERN.finditer(text))
    if not page_matches:
        # 无页码标记，整个文本视为一页
        pages.append({"page_num": 1, "content": text, "start": 0, "end": len(text)})
        for i in range(len(text)):
            char_to_page[i] = 1
        return pages, char_to_page

    # 有页码标记
    for i, match in enumerate(page_matches):
        content_start = match.end() + 1 if match.end() < len(text) else match.end()
        next_match_start = page_matches[i + 1].start() if i + 1 < len(page_matches) else len(text)
        page_content = text[content_start:next_match_start]
        page_num = int(match.group(1))

        pages.append({
            "page_num": page_num,
            "content": page_content,
            "start": content_start,
            "end": next_match_start,
        })

        # 建立字符位置到页码的映射
        for pos in range(content_start, next_match_start):
            char_to_page[pos] = page_num

    return pages, char_to_page


def _get_page_for_position(char_to_page: dict, pos: int) -> int | None:
    """根据字符位置查页码"""
    return char_to_page.get(pos)


def _extract_standard_info(text: str, filename: str = "") -> dict:
    """
    从文档文本中提取标准编号和标准名称。
    优先从前 2000 字符中查找。
    """
    head = text[:2000]
    result = {
        "standard_number": "",
        "standard_name": "",
    }

    # 1. 正则提取标准编号
    num_match = _STD_NUM_RE.search(head)
    if num_match:
        result["standard_number"] = num_match.group(0).strip()

    # 2. 标准名称：标准编号后面紧跟的一行或同一行后面的文本
    if result["standard_number"]:
        num_end = num_match.end()
        after_num = head[num_end:num_end + 200].strip()
        # 取到换行或句号
        name_end_chars = '\n\r。；'
        name_text = after_num
        for ch in name_end_chars:
            idx = after_num.find(ch)
            if idx > 0:
                name_text = after_num[:idx]
                break
        name_text = name_text.strip().lstrip('《').rstrip('》').strip()
        if len(name_text) >= 3:
            result["standard_name"] = name_text

    # 3. 备选：从"《》"中提取
    if not result["standard_name"]:
        bracket_match = re.search(r'《(.+?)》', head)
        if bracket_match:
            result["standard_name"] = bracket_match.group(1).strip()

    # 4. 用文件名兜底
    if not result["standard_name"] and filename:
        clean = filename
        for ext in ['.pdf', '.docx', '.txt', '.PDF', '.DOCX', '.TXT']:
            clean = clean.replace(ext, '')
        result["standard_name"] = clean.strip()

    return result


def _detect_major_sections(text: str) -> list[dict]:
    """
    检测文档的主要区域边界（增强版 v2 — 三层级联）。

    Layer 1: GB/T 标准格式（1 范围 / 2 规范性引用文件 / 3 术语和定义 / 前言 / 附录 / 参考文献）
    Layer 2: 通用标题格式（第X章 / X.X 标题 / 一、二、三...）
    Layer 3: 启发式短行检测（短行 + 后续长段落 → 可能是标题）

    返回区域列表，每个区域包含 type / title / start / end。
    """
    lines = text.split('\n')
    sections = []

    # ═══ Layer 1: GB/T 标准格式检测 ═══
    # 注：已移除 \b（中文后 \b 在 Python regex 中不匹配行尾换行符）
    header_patterns_l1 = [
        (r'^前\s*言\s*$', 'preface'),
        (r'^引\s*言\s*$', 'preface'),
        (r'^(?:1|一)[\s\.、]+\s*范围', 'scope'),
        (r'^(?:2|二)[\s\.、]+\s*规范性引用文件', 'references'),
        (r'^(?:3|三)[\s\.、]+\s*(?:术语和定义|术语与定义|定义)', 'terms'),
        (r'^附录\s*[A-Za-z]', 'appendix'),
        (r'^参考文献\s*$', 'references'),
    ]

    header_positions = []  # (line_index, type, title)

    for i, line in enumerate(lines):
        line_clean = line.strip()
        if not line_clean:
            continue
        # 跳过目录行
        if re.search(r'\.{5,}', line_clean) or re.search(r'…{2,}', line_clean):
            continue
        if re.match(r'^[\d\.\s]+$', line_clean):
            continue
        for pattern, stype in header_patterns_l1:
            if re.match(pattern, line_clean):
                header_positions.append((i, stype, line_clean))
                break

    # ═══ Layer 2: 通用标题检测（仅当 Layer 1 检测不足时触发）═══
    if len(header_positions) < 2:
        # 检测"第X章"、"第X节"
        generic_chapter = re.compile(r'^第[一二三四五六七八九十百\d]+[章节]')

        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean or len(line_clean) > 80:
                continue
            # 检查是否已被 Layer 1 捕获
            if any(hp[0] == i for hp in header_positions):
                continue
            # 模式1: 第X章 / 第X节
            if generic_chapter.match(line_clean):
                header_positions.append((i, 'clauses', line_clean))
                continue
            # 模式2: 编号标题（如 "1. 绪论"、"1.1 研究背景"、"一、概述"）
            if re.match(r'^\d+(?:\.\d+)*[\s\.、]+\S', line_clean) and len(line_clean) <= 50:
                # 确认不是普通数据行（避免将 "2024年3月" 误判）
                if not re.match(r'^\d{4}[\s 年]', line_clean):
                    header_positions.append((i, 'clauses', line_clean))
                    continue
            # 模式3: 中文序号标题（"一、..." "二、..."）
            if re.match(r'^[一二三四五六七八九十]+[、．.]\s*\S', line_clean) and len(line_clean) <= 50:
                header_positions.append((i, 'clauses', line_clean))
                continue

        # 按行号排序（Layer 2 可能乱序插入）
        header_positions.sort(key=lambda x: x[0])

    # ═══ Layer 3: 启发式短行检测（仅当仍不足时）═══
    if len(header_positions) < 2:
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
            if any(hp[0] == i for hp in header_positions):
                continue
            # 短行（5-20 字）且下一行是长段落（> 40 字）→ 可能是标题
            if 5 <= len(line_clean) <= 20:
                if i + 1 < len(lines) and len(lines[i + 1].strip()) > 40:
                    # 不包含句号等句子结束符
                    if not re.search(r'[。；，！？]', line_clean):
                        header_positions.append((i, 'clauses', line_clean))

    # ═══ 构建 sections ═══

    # 如果没检测到任何标题，整个文档视为 clauses
    if not header_positions:
        sections.append({
            "type": "clauses",
            "title": "正文",
            "start_line": 0,
            "end_line": len(lines),
            "content": text,
        })
        return sections

    # 第一个标题之前的内容 → 封面（有长度限制）
    first_header_line = header_positions[0][0]
    if first_header_line > 0:
        cover_text = '\n'.join(lines[:first_header_line]).strip()
        # 封面内容不应超过 3000 字（否则可能是漏检的长文档）
        if cover_text and 20 < len(cover_text) <= 3000:
            sections.append({
                "type": "cover",
                "title": "封面",
                "start_line": 0,
                "end_line": first_header_line,
                "content": cover_text,
            })
        elif cover_text and len(cover_text) > 3000:
            # 封面内容超长 → 截断封面，保留后续检测到的章节结构
            logger.warning(f"封面内容 {len(cover_text)} 字，超过 3000 字上限，截断处理")
            truncated_lines = lines[:first_header_line]
            # 取前 ~2000 字作为封面
            truncated = ""
            char_count = 0
            for line in truncated_lines:
                if char_count > 2000:
                    break
                truncated += line + "\n"
                char_count += len(line) + 1
            cover_text = truncated.strip()
            if cover_text and len(cover_text) > 20:
                sections.append({
                    "type": "cover",
                    "title": "封面(截断)",
                    "start_line": 0,
                    "end_line": first_header_line,
                    "content": cover_text,
                })

    # 按标题切分区域
    for idx, (line_idx, stype, title) in enumerate(header_positions):
        next_line = header_positions[idx + 1][0] if idx + 1 < len(header_positions) else len(lines)
        section_text = '\n'.join(lines[line_idx:next_line]).strip()

        char_start = sum(len(lines[k]) + 1 for k in range(line_idx))

        sections.append({
            "type": stype,
            "title": title,
            "start_line": line_idx,
            "end_line": next_line,
            "start_char": char_start,
            "content": section_text,
        })

    # 合并相邻同类区域
    merged = []
    for sec in sections:
        if merged and merged[-1]["type"] == sec["type"] == "preface":
            merged[-1]["end_line"] = sec["end_line"]
            merged[-1]["content"] += "\n" + sec["content"]
            merged[-1]["title"] = "前言/引言"
        elif merged and merged[-1]["type"] == sec["type"] == "references":
            merged[-1]["end_line"] = sec["end_line"]
            merged[-1]["content"] += "\n" + sec["content"]
        else:
            merged.append(sec)

    # 后处理：检测 terms 区域中是否包含正文条款
    processed = []
    for sec in merged:
        if sec.get("type") == "terms":
            sec_lines = sec["content"].split('\n')
            split_at = None
            for j, line in enumerate(sec_lines):
                line_s = line.strip()
                m = re.match(r'^(\d+)\s+\S', line_s)
                if m and m.group(1) not in ('1', '2', '3') and len(m.group(1)) <= 2:
                    split_at = j
                    break
            if split_at is not None and split_at > 0:
                terms_content = '\n'.join(sec_lines[:split_at]).strip()
                clauses_content = '\n'.join(sec_lines[split_at:]).strip()
                base_char = sec.get("start_char", 0)
                terms_char_len = sum(len(sec_lines[k]) + 1 for k in range(split_at))
                if terms_content:
                    processed.append({
                        "type": "terms",
                        "title": sec.get("title", "3 术语和定义"),
                        "start_line": sec.get("start_line", 0),
                        "end_line": sec.get("start_line", 0) + split_at,
                        "start_char": base_char,
                        "content": terms_content,
                    })
                if clauses_content:
                    processed.append({
                        "type": "clauses",
                        "title": "正文",
                        "start_line": sec.get("start_line", 0) + split_at,
                        "end_line": sec.get("end_line", 0),
                        "start_char": base_char + terms_char_len,
                        "content": clauses_content,
                    })
                continue
        processed.append(sec)

    return processed


def _parse_clause_tree(section_text: str) -> list[dict]:
    """
    从区域文本中解析条款层级树。
    返回嵌套的条款列表。
    """
    lines = section_text.split('\n')
    clauses = []
    stack = []  # (level, clause_dict)

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # 跳过区域标题行（如"3 术语和定义"）
        if any(re.match(p, line_stripped) for p in [
            r'^(?:1|一|2|二|3|三)[\s\.、]+',
            r'^前\s*言\s*$', r'^引\s*言\s*$',
        ]):
            continue

        # 匹配条款编号
        clause_match = _CLAUSE_LINE_RE.match(line_stripped)
        if not clause_match:
            # 非条款行，附加到当前最深条款的 content
            if stack:
                stack[-1][1]["content"] += "\n" + line_stripped
            elif clauses:
                clauses[-1]["content"] += "\n" + line_stripped
            continue

        number = clause_match.group(1)
        title = clause_match.group(2)
        level = len(number.split('.'))

        clause = {
            "number": number,
            "title": title,
            "content": line_stripped,
            "children": [],
            "level": level,
        }

        # 确定该条款在树中的位置
        while stack and stack[-1][0] >= level:
            stack.pop()

        if stack:
            stack[-1][1]["children"].append(clause)
        else:
            clauses.append(clause)

        stack.append((level, clause))

    return clauses


def _flatten_clauses(clauses: list[dict], parent_path: str = "") -> list[dict]:
    """
    将嵌套条款树展开为扁平的条款列表，每个带完整 section_path。
    """
    flat = []
    for c in clauses:
        path = (parent_path + " > " + c["number"] + " " + c["title"]) if parent_path else (c["number"] + " " + c["title"])
        flat.append({
            "number": c["number"],
            "title": c["title"],
            "content": c["content"],
            "section_path": path,
            "level": c["level"],
            "children": [ch["number"] for ch in c.get("children", [])],
        })
        if c.get("children"):
            flat.extend(_flatten_clauses(c["children"], path))
    return flat


def _extract_tables_from_text(text: str) -> list[dict]:
    """
    从文本中检测并提取表格区域。
    支持：
    1. 解析器生成的 [表格 N] 标记
    2. Markdown 表格（| ... | ... | 格式）
    3. 以"表"开头的标题段落
    """
    tables = []

    # 1. 检测 [表格 N] 标记
    for match in _TABLE_PATTERN.finditer(text):
        # 找到表格标记后的内容，直到下一个标记或文本结束
        start = match.end()
        # 找下一个结构标记
        next_marker = re.search(
            r'=====\s*\[(?:表格|文本框)\s*\d*\s*\]\s*=====|'
            r'=====\s*第\s*\d+\s*页\s*=====',
            text[start:]
        )
        end = start + next_marker.start() if next_marker else len(text)
        table_content = text[start:end].strip()

        if table_content:
            tables.append({
                "type": "table",
                "table_index": int(match.group(1)),
                "content": table_content,
                "start": match.start(),
                "end": end,
            })

    # 2. 检测 Markdown 表格（连续的 |...| 行）
    md_table_re = re.compile(r'(?:[|｜][^\n]*[|｜]\s*\n){2,}', re.MULTILINE)
    for match in md_table_re.finditer(text):
        # 避免和已检测的 [表格 N] 区域重叠
        pos = match.start()
        if any(t["start"] <= pos < t["end"] for t in tables):
            continue
        tables.append({
            "type": "table",
            "table_index": None,
            "content": match.group(0).strip(),
            "start": pos,
            "end": match.end(),
        })

    # 3. 检测"表N"标题（后续文本直到下一个条款或空行）
    for match in _TABLE_TITLE_RE.finditer(text):
        pos = match.start()
        # 避免重叠
        if any(t["start"] <= pos < t["end"] for t in tables):
            continue
        # 取标题行 + 后续内容直到条款边界或两个空行
        content_start = match.start()
        content_end = match.end()
        after = text[content_end:]
        # 找下一个条款编号或两个连续空行
        boundary = re.search(r'\n\d+(?:\.\d+)*\s+\S|\n\s*\n\s*\n', after)
        if boundary:
            content_end += boundary.start()
        else:
            content_end += min(800, len(after))  # 最多取800字
        tables.append({
            "type": "table",
            "table_index": None,
            "content": text[content_start:content_end].strip(),
            "start": content_start,
            "end": content_end,
        })

    return tables


def _extract_figures_from_text(text: str) -> list[dict]:
    """检测图示/公式区域"""
    figures = []
    for match in _FIGURE_TITLE_RE.finditer(text):
        pos = match.start()
        content_start = match.start()
        content_end = match.end()
        after = text[content_end:]
        boundary = re.search(r'\n\d+(?:\.\d+)*\s+\S|\n\s*\n\s*\n', after)
        if boundary:
            content_end += boundary.start()
        else:
            content_end += min(500, len(after))
        figures.append({
            "type": "figure",
            "content": text[content_start:content_end].strip(),
            "start": content_start,
            "end": content_end,
        })
    return figures


def parse_document_structure(text: str, filename: str = "") -> dict:
    """
    主入口：解析文档全文，返回结构化数据。
    """
    if not text or text.strip() == "":
        return {"sections": [], "standard_info": {}, "tables": [], "figures": []}

    # 1. 页码映射
    pages, char_to_page = _map_page_numbers(text)

    # 2. 提取标准信息
    standard_info = _extract_standard_info(text, filename)

    # 3. 检测主区域
    sections = _detect_major_sections(text)

    # 4. 对每个区域解析条款树
    parsed_sections = []
    for sec in sections:
        parsed = dict(sec)

        if sec["type"] in ("clauses",):
            parsed["clauses"] = _parse_clause_tree(sec["content"])
            parsed["flat_clauses"] = _flatten_clauses(parsed["clauses"])

        elif sec["type"] == "terms":
            parsed["clauses"] = _parse_clause_tree(sec["content"])
            parsed["flat_clauses"] = _flatten_clauses(parsed["clauses"])

        elif sec["type"] == "appendix":
            parsed["clauses"] = _parse_clause_tree(sec["content"])
            # 修正路径加"附录"前缀
            appendix_label = sec.get("title", "附录")
            parsed["flat_clauses"] = _flatten_clauses(parsed["clauses"], appendix_label)

        elif sec["type"] == "references":
            # 尝试提取每个引用标准
            refs = []
            for ref_match in _STD_NUM_RE.finditer(sec["content"]):
                refs.append(ref_match.group(0).strip())
            parsed["individual_refs"] = refs

        parsed_sections.append(parsed)

    # 5. 检测全文中的表格和图
    tables = _extract_tables_from_text(text)
    figures = _extract_figures_from_text(text)

    return {
        "sections": parsed_sections,
        "standard_info": standard_info,
        "tables": tables,
        "figures": figures,
        "pages": pages,
        "char_to_page": char_to_page,
    }


# ═══════════════════════════════════════════════════════════════
# v2 智能切片 — 分区域策略
# ═══════════════════════════════════════════════════════════════

def _build_chunk_prefix(standard_info: dict, section_path: str = "", page_number: int | None = None) -> str:
    """构建 chunk 的前缀元信息"""
    parts = []
    std_num = standard_info.get("standard_number", "")
    std_name = standard_info.get("standard_name", "")
    if std_num and std_name:
        parts.append(f"标准：{std_name}（{std_num}）")
    elif std_num:
        parts.append(f"标准编号：{std_num}")
    elif std_name:
        parts.append(f"标准：{std_name}")

    if section_path:
        parts.append(f"章节路径：{section_path}")

    if page_number:
        parts.append(f"页码：第 {page_number} 页")

    if not parts:
        return ""
    return "\n".join(parts) + "\n---\n"


def _chunk_cover(sec: dict, std_info: dict, char_to_page: dict) -> list[dict]:
    """
    封面 → 1 个 chunk（含标准编号/名称）。
    如果 section 内容超过 2000 字，说明章节检测失败导致正文被误归入封面，
    此时回退到段落级通用切片。
    """
    content = sec.get("content", "")
    if not content or len(content.strip()) < 20:
        return []

    # 内容量守卫：封面不应超过 2000 字
    if len(content) > 2000:
        logger.warning(f"封面内容过长 ({len(content)} 字)，回退到段落切片")
        return _chunk_fallback_paragraphs(content, std_info, char_to_page, "cover")

    # 控制在 500 字以内
    text = content[:500].strip()
    char_pos = sec.get("start_char", 0)
    page = _get_page_for_position(char_to_page, char_pos)

    prefix = _build_chunk_prefix(std_info, "封面", page)
    return [{
        "content": prefix + text,
        "chunk_type": "cover",
        "section_path": "封面",
        "section_number": "",
        "page_number": page,
    }]


def _chunk_preface(sec: dict, std_info: dict, char_to_page: dict) -> list[dict]:
    """前言 → 1 个 chunk，超过1200字按段落拆分"""
    content = sec.get("content", "")
    if not content:
        return []
    char_pos = sec.get("start_char", 0)
    page = _get_page_for_position(char_to_page, char_pos)

    if len(content) <= 1200:
        prefix = _build_chunk_prefix(std_info, "前言", page)
        return [{
            "content": prefix + content.strip(),
            "chunk_type": "preface",
            "section_path": "前言",
            "section_number": "",
            "page_number": page,
        }]

    # 长前言按段落拆分
    chunks = []
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
    i = 0
    while i < len(paragraphs):
        merged = paragraphs[i]
        j = i + 1
        while j < len(paragraphs) and len(merged) + len(paragraphs[j]) < 800:
            merged += "\n\n" + paragraphs[j]
            j += 1
        prefix = _build_chunk_prefix(std_info, f"前言({i + 1}/{len(paragraphs)})", page)
        chunks.append({
            "content": prefix + merged,
            "chunk_type": "preface",
            "section_path": "前言",
            "section_number": "",
            "page_number": page,
        })
        i = j
    return chunks


def _chunk_scope(sec: dict, std_info: dict, char_to_page: dict) -> list[dict]:
    """范围 → 1 个 chunk"""
    content = sec.get("content", "")
    if not content:
        return []
    char_pos = sec.get("start_char", 0)
    page = _get_page_for_position(char_to_page, char_pos)
    prefix = _build_chunk_prefix(std_info, sec.get("title", "1 范围"), page)
    return [{
        "content": prefix + content.strip(),
        "chunk_type": "scope",
        "section_path": sec.get("title", "1 范围"),
        "section_number": "1",
        "page_number": page,
    }]


def _chunk_references(sec: dict, std_info: dict, char_to_page: dict) -> list[dict]:
    """规范性引用文件 → 1 个 chunk（或按引用分条）"""
    content = sec.get("content", "")
    if not content:
        return []
    char_pos = sec.get("start_char", 0)
    page = _get_page_for_position(char_to_page, char_pos)
    prefix = _build_chunk_prefix(std_info, sec.get("title", "2 规范性引用文件"), page)
    # 如果文本较短（< 1000字），一个 chunk
    if len(content) <= 1000:
        return [{
            "content": prefix + content.strip(),
            "chunk_type": "references",
            "section_path": sec.get("title", "2 规范性引用文件"),
            "section_number": "2",
            "page_number": _get_page_for_position(char_to_page, sec.get("start_line", 0)),
        }]

    # 较长则拆分（按引用标准条目）
    chunks = []
    refs = sec.get("individual_refs", [])
    if refs and len(refs) > 5:
        # 每个引用标准单独 chunk
        ref_lines = content.split('\n')
        current_ref = []
        current_idx = 0
        for line in ref_lines:
            current_ref.append(line)
            ref_text = '\n'.join(current_ref)
            if len(ref_text) > 200 or any(r in line for r in refs[len(chunks):]):
                chunk_prefix = _build_chunk_prefix(
                    std_info,
                    f"2 规范性引用文件 > {refs[current_idx] if current_idx < len(refs) else ''}",
                    _get_page_for_position(char_to_page, sec.get("start_line", 0))
                )
                chunks.append({
                    "content": chunk_prefix + ref_text.strip(),
                    "chunk_type": "references",
                    "section_path": f"2 规范性引用文件 > {refs[current_idx] if current_idx < len(refs) else ''}",
                    "section_number": "2",
                    "page_number": _get_page_for_position(char_to_page, sec.get("start_line", 0)),
                })
                current_ref = []
                current_idx += 1
        if current_ref:
            chunk_prefix = _build_chunk_prefix(std_info, "2 规范性引用文件",
                                                _get_page_for_position(char_to_page, sec.get("start_line", 0)))
            chunks.append({
                "content": chunk_prefix + '\n'.join(current_ref).strip(),
                "chunk_type": "references",
                "section_path": "2 规范性引用文件",
                "section_number": "2",
                "page_number": _get_page_for_position(char_to_page, sec.get("start_line", 0)),
            })
    else:
        chunks.append({
            "content": prefix + content.strip(),
            "chunk_type": "references",
            "section_path": sec.get("title", "2 规范性引用文件"),
            "section_number": "2",
            "page_number": _get_page_for_position(char_to_page, sec.get("start_line", 0)),
        })

    return chunks


def _chunk_terms(sec: dict, std_info: dict, char_to_page: dict) -> list[dict]:
    """术语和定义 → 每条术语 1 个 chunk，短术语合并"""
    flat = sec.get("flat_clauses", [])
    char_pos = sec.get("start_char", 0)
    page = _get_page_for_position(char_to_page, char_pos)

    if not flat:
        content = sec.get("content", "")
        if not content:
            return []
        prefix = _build_chunk_prefix(std_info, sec.get("title", "3 术语和定义"), page)
        return [{
            "content": prefix + content.strip(),
            "chunk_type": "term",
            "section_path": sec.get("title", "3 术语和定义"),
            "section_number": "3",
            "page_number": page,
        }]

    chunks = []
    buffer = []
    buffer_len = 0

    def flush_buffer():
        if not buffer:
            return
        merged_content = "\n".join(b["content"] for b in buffer)
        nums = [b["number"] for b in buffer]
        path = f"3 术语和定义 > {nums[0]}"
        if len(nums) > 1:
            path += f" - {nums[-1]}"
        chunk_prefix = _build_chunk_prefix(std_info, path, page)
        chunks.append({
            "content": chunk_prefix + merged_content,
            "chunk_type": "term",
            "section_path": path,
            "section_number": nums[0],
            "page_number": page,
        })

    for term in flat:
        term_len = len(term["content"])
        if buffer_len + term_len > 500 and buffer:
            flush_buffer()
            buffer = []
            buffer_len = 0
        buffer.append(term)
        buffer_len += term_len

    flush_buffer()
    return chunks


def _chunk_clauses(flat_clauses: list[dict], std_info: dict, char_to_page: dict, sec_start_char: int) -> list[dict]:
    """
    正文条款切片：按条款号切分，300-500字基准。
    - 短条款（< 300字）：合并相邻同级条款，不跨父级主题
    - 适中条款（300-500字）：独立成 chunk
    - 长条款（> 1200字）：按自然段继续拆分，保留父级路径
    """
    if not flat_clauses:
        return []

    chunks = []
    buffer = []        # 短条款合并缓冲区
    buffer_len = 0
    buffer_paths = []
    page = _get_page_for_position(char_to_page, sec_start_char)

    def flush():
        if not buffer:
            return
        merged_content = "\n".join(b["content"] for b in buffer)
        paths = [b["section_path"] for b in buffer]
        path = paths[0] if len(paths) == 1 else paths[0].rsplit(" > ", 1)[0] + f" > {buffer[0]['number']} - {buffer[-1]['number']}"
        chunk_prefix = _build_chunk_prefix(std_info, path, page)
        chunks.append({
            "content": chunk_prefix + merged_content,
            "chunk_type": "clause",
            "section_path": path,
            "section_number": buffer[0]["number"],
            "page_number": page,
        })

    def add_long_clause(clause: dict):
        """长条款拆分"""
        content = clause["content"]
        if len(content) <= 1200:
            chunk_prefix = _build_chunk_prefix(std_info, clause["section_path"], page)
            chunks.append({
                "content": chunk_prefix + content,
                "chunk_type": "clause",
                "section_path": clause["section_path"],
                "section_number": clause["number"],
                "page_number": page,
            })
            return

        # 超过1200字，按自然段拆分
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        if len(paragraphs) <= 1:
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]

        part = 1
        total = max(1, len(paragraphs) // max(1, (len(content) // 500)))
        total = max(1, min(total, len(paragraphs)))

        i = 0
        while i < len(paragraphs):
            merged = paragraphs[i]
            j = i + 1
            while j < len(paragraphs) and len(merged) + len(paragraphs[j]) < 500:
                merged += "\n\n" + paragraphs[j]
                j += 1
            sub_path = f"{clause['section_path']}({part}/{total})" if total > 1 else clause["section_path"]
            chunk_prefix = _build_chunk_prefix(std_info, sub_path, page)
            chunks.append({
                "content": chunk_prefix + merged,
                "chunk_type": "clause",
                "section_path": sub_path,
                "section_number": clause["number"],
                "page_number": page,
            })
            part += 1
            i = j

    for clause in flat_clauses:
        clause_len = len(clause["content"])

        if clause_len < 300:
            # 短条款 → 合并缓冲区
            if buffer and buffer[0]["level"] != clause["level"]:
                # 层级变化，先 flush
                flush()
                buffer = []
                buffer_len = 0
                buffer_paths = []
            if buffer_len + clause_len > 600:
                flush()
                buffer = []
                buffer_len = 0
                buffer_paths = []
            buffer.append(clause)
            buffer_len += clause_len
            buffer_paths.append(clause["section_path"])
        else:
            # 先 flush 缓冲区
            flush()
            buffer = []
            buffer_len = 0
            buffer_paths = []
            # 处理当前条款
            add_long_clause(clause)

    flush()
    return chunks


def _chunk_tables(tables: list[dict], std_info: dict, char_to_page: dict) -> list[dict]:
    """表格 → 每个表格 1 个 chunk，保持完整"""
    chunks = []
    for t in tables:
        content = t.get("content", "")
        if not content:
            continue
        page = _get_page_for_position(char_to_page, t.get("start", 0))
        table_label = f"表{t['table_index']}" if t.get("table_index") else "表格"
        chunk_prefix = _build_chunk_prefix(std_info, table_label, page)
        chunks.append({
            "content": chunk_prefix + content,
            "chunk_type": "table",
            "section_path": table_label,
            "section_number": "",
            "page_number": page,
        })
    return chunks


def _chunk_figures(figures: list[dict], std_info: dict, char_to_page: dict) -> list[dict]:
    """图示/公式 → 每图 1 个 chunk，图题+说明+上下文"""
    chunks = []
    for f in figures:
        content = f.get("content", "")
        if not content:
            continue
        page = _get_page_for_position(char_to_page, f.get("start", 0))
        chunk_prefix = _build_chunk_prefix(std_info, "图示/公式", page)
        chunks.append({
            "content": chunk_prefix + content,
            "chunk_type": "figure",
            "section_path": "图示/公式",
            "section_number": "",
            "page_number": page,
        })
    return chunks


def _chunk_appendix(sec: dict, std_info: dict, char_to_page: dict) -> list[dict]:
    """附录 → 按附录内条款切分，与正文条款逻辑一致"""
    flat = sec.get("flat_clauses", [])
    char_pos = sec.get("start_char", 0)
    page = _get_page_for_position(char_to_page, char_pos)

    if not flat:
        content = sec.get("content", "")
        if not content:
            return []
        path = sec.get("title", "附录")
        prefix = _build_chunk_prefix(std_info, path, page)
        return [{
            "content": prefix + content.strip(),
            "chunk_type": "appendix",
            "section_path": path,
            "section_number": "",
            "page_number": page,
        }]

    return _chunk_clauses(flat, std_info, char_to_page, char_pos)


def _chunk_fallback_paragraphs(
    text: str,
    std_info: dict,
    char_to_page: dict,
    section_label: str = "正文",
    target_size: int = 500
) -> list[dict]:
    """
    通用段落级兜底切片。
    按空行切段，相邻短段合并到 target_size 左右。
    用于章节检测失败时的保障机制，确保文档至少被合理分割。
    """
    if not text or text.strip() == "":
        return []

    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    if not paragraphs:
        # 无空行 → 按单行切
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]

    if not paragraphs:
        return []

    chunks = []
    i = 0
    part_num = 1
    total_estimate = max(1, len(text) // target_size)

    while i < len(paragraphs):
        merged = paragraphs[i]
        j = i + 1
        while j < len(paragraphs) and len(merged) + len(paragraphs[j]) < target_size:
            merged += "\n\n" + paragraphs[j]
            j += 1

        # 估算页码
        char_pos = sum(len(p) + 2 for p in paragraphs[:i])
        page = _get_page_for_position(char_to_page, char_pos)

        path = f"{section_label}（{part_num}/{total_estimate}）" if total_estimate > 1 else section_label
        prefix = _build_chunk_prefix(std_info, path, page)
        chunks.append({
            "content": prefix + merged,
            "chunk_type": "clause",
            "section_path": path,
            "section_number": "",
            "page_number": page,
        })
        part_num += 1
        i = j

    return chunks


def _classify_document_type(text: str, filename: str = "") -> str:
    """
    文档类型分类器。
    检查文本前 3000 字 + 文件名，返回类型标签。

    返回:
      national_standard  — GB, GB/T, GBZ, GJB
      industry_standard  — NY/T, HJ/T, SN/T, YY/T, DBxx/T 等行业标准
      enterprise_standard — Q/ 企业标准
      generic            — 通用文档
    """
    head = text[:3000] if len(text) > 3000 else text
    std_info = _extract_standard_info(text, filename)
    std_num = std_info.get("standard_number", "")

    # ── 按标准编号前缀判定 ──
    # 国家标准
    if re.search(r'\b(?:GB[ /TZ]|GJB)\s*\d', head, re.IGNORECASE):
        return "national_standard"

    # 行业标准
    industry_prefixes = (
        r'NY[ /T]|HJ[ /T]|SN[ /T]|YY[ /T]|QB[ /T]|'
        r'DB\d+[ /T]|CJJ|JGJ|GA[ /]|SY[ /T]|DL[ /T]|'
        r'WS[ /T]|MH[ /T]|JT[ /T]|LS[ /T]|LY[ /T]|SC[ /T]|SB[ /T]|'
        r'CY[ /T]|JB[ /T]|SJ[ /T]|CB[ /T]|CH[ /T]|CJ[ /T]|'
        r'GY[ /T]|HY[ /T]|JC[ /T]|JR[ /T]|NB[ /T]|'
        r'QC[ /T]|QX[ /T]|SH[ /T]|SL[ /T]|'
        r'TY[ /T]|WB[ /T]|WH[ /T]|YC[ /T]|YZ[ /T]|ZY[ /T]'
    )
    if re.search(r'\b(?:' + industry_prefixes + r')\s*\d', head, re.IGNORECASE):
        return "industry_standard"

    # 企业标准
    if re.search(r'\bQ/\s*\w', head):
        return "enterprise_standard"

    # ── 有编号结构但无标准编号 → 按行业标准处理 ──
    if re.search(r'\n\d+(?:\.\d+)*\s+\S', head):
        return "industry_standard"

    return "generic"


def smart_split_v2(text: str, filename: str = "") -> list[dict]:
    """
    v3 智能文档切片主入口 — 按文档类型分流策略。
    - national_standard → GB/T 条款层级切片（严格）
    - industry_standard → 行业标准切片（灵活条款）
    - enterprise_standard → 企业标准切片（灵活编号）
    - generic → 语义段落切片（兜底）

    语义保障：每个 chunk 以完整段落/条款为最小单位，绝不在句子中间截断。
    """
    if not text or text.strip() == "":
        return []

    doc_type = _classify_document_type(text, filename)

    # ── national_standard / industry_standard ──
    if doc_type in ("national_standard", "industry_standard"):
        struct = parse_document_structure(text, filename)
        std_info = struct.get("standard_info", {})
        char_to_page = struct.get("char_to_page", {})

        all_chunks = []
        for sec in struct.get("sections", []):
            stype = sec.get("type", "")
            if stype == "cover":
                all_chunks.extend(_chunk_cover(sec, std_info, char_to_page))
            elif stype == "preface":
                all_chunks.extend(_chunk_preface(sec, std_info, char_to_page))
            elif stype == "scope":
                all_chunks.extend(_chunk_scope(sec, std_info, char_to_page))
            elif stype == "references":
                all_chunks.extend(_chunk_references(sec, std_info, char_to_page))
            elif stype == "terms":
                all_chunks.extend(_chunk_terms(sec, std_info, char_to_page))
            elif stype == "clauses":
                flat = sec.get("flat_clauses", [])
                char_pos = sec.get("start_char", 0)
                if flat:
                    all_chunks.extend(_chunk_clauses(flat, std_info, char_to_page, char_pos))
                else:
                    content = sec.get("content", "")
                    if content:
                        all_chunks.extend(
                            _chunk_fallback_paragraphs(content, std_info, char_to_page,
                                                       section_label=sec.get("title", "正文"))
                        )
            elif stype == "appendix":
                all_chunks.extend(_chunk_appendix(sec, std_info, char_to_page))

        all_chunks.extend(_chunk_tables(struct.get("tables", []), std_info, char_to_page))
        all_chunks.extend(_chunk_figures(struct.get("figures", []), std_info, char_to_page))

        # 兜底保障
        if not all_chunks:
            all_chunks = _chunk_fallback_paragraphs(text, std_info, char_to_page, "全文")
        if len(text.strip()) > 2000 and len(all_chunks) < 3:
            all_chunks = _chunk_fallback_paragraphs(text, std_info, char_to_page, "全文")

        _assign_parent_refs(all_chunks)
        return all_chunks

    # ── enterprise_standard ──
    if doc_type == "enterprise_standard":
        pages, char_to_page = _map_page_numbers(text)
        std_info = _extract_standard_info(text, filename)
        sections = _detect_major_sections(text)

        all_chunks = []
        for sec in sections:
            stype = sec.get("type", "")
            content = sec.get("content", "")
            if not content:
                continue
            if stype == "cover":
                all_chunks.extend(_chunk_cover(sec, std_info, char_to_page))
            elif stype == "preface":
                all_chunks.extend(_chunk_preface(sec, std_info, char_to_page))
            elif stype in ("clauses", "scope", "references", "terms", "appendix"):
                flat = sec.get("flat_clauses", []) if "flat_clauses" in sec else _flatten_clauses(_parse_clause_tree(content))
                char_pos = sec.get("start_char", 0)
                if flat:
                    all_chunks.extend(_chunk_clauses(flat, std_info, char_to_page, char_pos))
                else:
                    all_chunks.extend(_chunk_fallback_paragraphs(content, std_info, char_to_page, sec.get("title", "正文")))
            else:
                all_chunks.extend(_chunk_fallback_paragraphs(content, std_info, char_to_page, sec.get("title", "正文")))

        tables = _extract_tables_from_text(text)
        figures = _extract_figures_from_text(text)
        all_chunks.extend(_chunk_tables(tables, std_info, char_to_page))
        all_chunks.extend(_chunk_figures(figures, std_info, char_to_page))

        if not all_chunks:
            all_chunks = _chunk_fallback_paragraphs(text, std_info, char_to_page, "全文")
        if len(text.strip()) > 2000 and len(all_chunks) < 3:
            all_chunks = _chunk_fallback_paragraphs(text, std_info, char_to_page, "全文")

        _assign_parent_refs(all_chunks)
        return all_chunks

    # ── generic ──
    pages, char_to_page = _map_page_numbers(text)
    std_info = _extract_standard_info(text, filename)
    all_chunks = _chunk_fallback_paragraphs(text, std_info, char_to_page, "全文")

    return all_chunks


def _assign_parent_refs(chunks: list[dict]):
    """
    根据 section_path 和 section_number 为 chunks 建立父子引用。
    父级 = 路径中包含当前路径前缀、section_number 是当前父级编号的 chunk。
    用 parent_section_number 和 parent_chunk_index 字段标记。
    """
    n = len(chunks)
    for i in range(n):
        cur = chunks[i]
        cur_num = cur.get("section_number", "")
        cur_path = cur.get("section_path", "")

        if not cur_num:
            continue

        # 找父级编号：去掉最后一级（如 "5.2.1" → "5.2"）
        parts = cur_num.rsplit(".", 1)
        if len(parts) != 2:
            continue
        parent_num = parts[0]

        # 向前查找最近的同路径前缀 chunk
        for j in range(i - 1, -1, -1):
            prev = chunks[j]
            prev_num = prev.get("section_number", "")
            prev_path = prev.get("section_path", "")
            if prev_num == parent_num and cur_path.startswith(prev_path.split("(")[0]):
                cur["parent_section_number"] = parent_num
                cur["parent_chunk_index"] = j
                break


# ═══════════════════════════════════════════════════════════════
# 切片存储
# ═══════════════════════════════════════════════════════════════

def save_document_chunk(db, chunks, document_id: int):
    save_chunk = []
    try:
        for index, chunk in enumerate(chunks):
            # 兼容 v1 (str) 和 v2 (dict)
            if isinstance(chunk, str):
                content = chunk
                chunk_type = None
                section_path = None
                section_number = None
                parent_chunk_id = None
                page_number = None
            else:
                content = chunk.get("content", "")
                chunk_type = chunk.get("chunk_type")
                section_path = chunk.get("section_path")
                section_number = chunk.get("section_number")
                page_number = chunk.get("page_number")
                # parent_chunk_index 是临时索引，需要在全部保存后换算为实际 ID
                parent_chunk_id = None  # 第二轮赋值

            embedding = services.embedding_service.generate_embedding(content)
            embedding_json = services.embedding_service.embedding_to_json(embedding)
            logger.debug(f"chunk #{index} doc={document_id} len={len(content)} type={chunk_type} path={section_path}")
            chunk_record = DocumentChunk(
                document_id=document_id,
                chunk_index=index,
                content=content,
                embedding=embedding_json,
                chunk_type=chunk_type,
                section_path=section_path,
                section_number=section_number,
                parent_chunk_id=parent_chunk_id,
                page_number=page_number,
            )
            db.add(chunk_record)
            save_chunk.append(chunk_record)
        db.flush()  # 获取 ID

        # 第二轮：建立父子关系（parent_chunk_index → parent_chunk_id）
        for idx, chunk_data in enumerate(chunks):
            if isinstance(chunk_data, dict) and chunk_data.get("parent_chunk_index") is not None:
                parent_idx = chunk_data["parent_chunk_index"]
                if 0 <= parent_idx < len(save_chunk):
                    save_chunk[idx].parent_chunk_id = save_chunk[parent_idx].id
                    # 写入 chunk_relations 表
                    from models import ChunkRelation
                    db.add(ChunkRelation(
                        chunk_id=save_chunk[idx].id,
                        related_chunk_id=save_chunk[parent_idx].id,
                        relation_type="parent"
                    ))
                    db.add(ChunkRelation(
                        chunk_id=save_chunk[parent_idx].id,
                        related_chunk_id=save_chunk[idx].id,
                        relation_type="child"
                    ))

        db.commit()
    except SQLAlchemyError as e:
        logger.error(f"保存chunks 或 embedding 失败：{e!r}")
        traceback.print_exc()
        db.rollback()
        return None
    return save_chunk


def delete_chunks_by_id(db, document_id: int):
    chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).all()
    try:
        for chunk in chunks:
            db.delete(chunk)
            db.commit()
    except SQLAlchemyError:
        db.rollback()
        return False
    return True


# ═══════════════════════════════════════════════════════════════
# 检索函数
# ═══════════════════════════════════════════════════════════════

def semantic_search_chunks(db, question: str, limit: int = 5, document_ids: list[int] | None = None):
    """语义检索：支持多标准。模型不可用时降级返回空，不阻塞关键词检索。"""
    try:
        question_embedding = services.embedding_service.generate_query_embedding(question)
    except Exception as e:
        logger.warning(f"语义检索跳过（embedding 模型不可用）: {e}")
        return []

    query = db.query(DocumentChunk).filter(
        DocumentChunk.embedding.isnot(None)
    )
    if document_ids:
        query = query.filter(DocumentChunk.document_id.in_(document_ids))
    chunks = query.all()
    scored_chunks = []

    for chunk in chunks:
        chunk_embedding = services.embedding_service.json_to_embedding(chunk.embedding)
        score = services.embedding_service.cosine_similarity(
            question_embedding,
            chunk_embedding
        )
        if score < 0.55:
            continue
        scored_chunks.append({
            "chunk": chunk,
            "score": score
        })

    scored_chunks.sort(key=lambda item: item["score"], reverse=True)
    return scored_chunks[:limit]


# ═══════════════════════════════════════════════════════════════
# 智能搜索词提取 (jieba 分词 + 条款检测)
# ═══════════════════════════════════════════════════════════════

# 常见条款编号模式
_CLAUSE_PATTERN = re.compile(
    r'(?:第[一二三四五六七八九十百\d]+[条款章节])|'      # 第3条、第一章
    r'(?:\d+\.\d+(?:\.\d+)?)'                           # 3.2、5.1.3
)


def _extract_clause_refs(question: str) -> list[str]:
    """从问题中提取条款编号引用，如 '第3.2条'、'5.1'"""
    matches = _CLAUSE_PATTERN.findall(question)
    return list(dict.fromkeys(matches))  # 去重保序


def _jieba_extract_keywords(question: str) -> list[str]:
    """
    使用 jieba 分词 + 词性标注提取有意义的搜索关键词。
    保留:
    - 名词 (n, nr, ns, nt, nz)
    - 动词 (v, vn)
    - 英文/数字 (eng)
    - 专用术语
    """
    # 先做一遍精确模式分词
    words = jieba.lcut(question)
    # 再做词性标注
    pos_words = pseg.cut(question)

    keywords = []

    # 停用词表（轻度过滤，不过度删除）
    stopwords = {
        "请问", "请", "一下", "这个", "那个",
        "的", "是", "吗", "呢", "啊", "吧",
        "是什么", "有哪些", "哪些", "多少",
        "怎么", "如何", "什么",
        "一个", "一些", "这个", "那个", "哪个",
    }

    for word, flag in pos_words:
        word = word.strip()
        if not word or len(word) < 2:
            continue
        if word in stopwords:
            continue

        # 保留名词类（含专有名词、地名、机构名等）
        if flag.startswith('n'):  # n, nr, ns, nt, nz, ng, nl
            keywords.append(word)
        # 保留动词
        elif flag.startswith('v'):  # v, vn, vd
            keywords.append(word)
        # 保留英文和数字串
        elif flag == 'eng':
            keywords.append(word)

    # 额外：从精确模式分出的词中提取未覆盖的词（作为补充）
    for word in words:
        word = word.strip()
        if len(word) >= 2 and word not in stopwords and word not in keywords:
            # 保留未被词性标注覆盖但可能在精确模式中的词
            if any('一' <= ch <= '鿿' for ch in word):
                keywords.append(word)

    return keywords


def extract_search_terms(question: str) -> list[str]:
    """
    从用户问题中智能提取搜索词，结合：
    1. jieba 分词 + 词性过滤
    2. 条款编号检测
    3. 意图检测扩展
    返回去重的搜索词列表
    """
    final_terms = []

    # 1. jieba 分词提取关键词
    jieba_keywords = _jieba_extract_keywords(question)
    final_terms.extend(jieba_keywords)

    # 2. 条款编号检测
    clause_refs = _extract_clause_refs(question)
    final_terms.extend(clause_refs)

    # 3. 意图检测扩展（只加精准短语，不加单字通用词）
    intent = infer_question_intent(question)
    if intent == "scope":
        final_terms.extend(["适用范围", "适用于", "本文件适用于", "本标准适用于"])
    elif intent == "definition":
        final_terms.extend(["术语", "定义"])
    elif intent == "requirement":
        final_terms.extend(["要求", "规定", "指标"])

    # 4. 过滤：去掉太宽泛的通用词（匹配太多 chunk 的词）
    _BROAD_TERMS = {
        "场景", "哪些", "什么", "怎么", "如何",
        "这个", "那个", "文档", "文件",
        "请问", "一下", "多少", "是什么", "有哪些",
    }

    # 去重 + 过滤
    seen = set()
    deduped = []
    for term in final_terms:
        term = term.strip()
        if not term:
            continue
        if len(term) < 2:  # 跳过单字
            continue
        if term in _BROAD_TERMS:
            continue
        if term not in seen:
            deduped.append(term)
            seen.add(term)

    return deduped


def extract_keyword_terms(db, terms: list[str], limit: int = 5, document_ids: list[int] | None = None):
    """多关键词 OR 检索，支持多标准"""
    conditions = []
    for term in terms:
        if term and term.strip():
            conditions.append(DocumentChunk.content.contains(term.strip()))

    if not conditions:
        return []

    query = db.query(DocumentChunk).filter(or_(*conditions))
    if document_ids:
        query = query.filter(DocumentChunk.document_id.in_(document_ids))
    return query.limit(limit).all()


def _match_question_to_documents(db, question: str) -> int | None:
    """
    检测问题中是否提到了某个已上传文档的文件名。
    如果是，自动返回其 document_id，用于限定检索范围。
    纯字符串匹配，零延迟，不调用 LLM。
    """
    if not question or len(question) < 3:
        return None

    docs = db.query(Document).all()
    best_match = None
    best_len = 0

    for doc in docs:
        filename = doc.filename or ""
        # 去掉扩展名，得到纯文档名
        filename_clean = filename
        for ext in ['.pdf', '.docx', '.txt', '.PDF', '.DOCX', '.TXT']:
            filename_clean = filename_clean.replace(ext, '')

        filename_clean = filename_clean.strip()
        if len(filename_clean) < 3:
            continue

        # 完全匹配（文档名完整出现在问题中）
        if filename_clean in question:
            if len(filename_clean) > best_len:
                best_match = doc.id
                best_len = len(filename_clean)

        # 部分匹配：文档名中的关键词（长度 ≥ 4 的词）出现在问题中
        if best_match is None and len(filename_clean) >= 4:
            # 用 jieba 分词后逐词匹配
            words = jieba.lcut(filename_clean)
            meaningful_words = [w for w in words if len(w) >= 3]
            if meaningful_words and all(w in question for w in meaningful_words[:3]):
                best_match = doc.id
                best_len = sum(len(w) for w in meaningful_words[:3])

    return best_match


# ═══════════════════════════════════════════════════════════════
# IDF 缓存 + 查询扩展
# ═══════════════════════════════════════════════════════════════

_idf_cache: dict[str, float] | None = None
_idf_cache_chunk_count: int = 0


def _compute_idf_weights(db, force_refresh: bool = False) -> dict[str, float]:
    """
    预计算所有 chunk 中词条的 IDF 权重。
    采用平滑 IDF: log((N + 1) / (df + 1)) + 1
    结果缓存在 _idf_cache 中，仅当 chunk 数量变化时刷新。
    """
    global _idf_cache, _idf_cache_chunk_count

    total_chunks = db.query(DocumentChunk).count()
    if not force_refresh and _idf_cache is not None and _idf_cache_chunk_count == total_chunks:
        return _idf_cache

    logger.info(f"重新计算 IDF 权重 ({total_chunks} chunks)...")

    # 遍历所有 chunk 建立文档频率
    doc_freq: dict[str, int] = {}
    chunks = db.query(DocumentChunk).all()

    for chunk in chunks:
        content = chunk.content or ""
        # 用 jieba 分词
        words = set(jieba.lcut(content))
        for w in words:
            w = w.strip()
            if len(w) >= 2:
                doc_freq[w] = doc_freq.get(w, 0) + 1

    N = total_chunks if total_chunks > 0 else 1
    import math
    idf = {}
    for word, df in doc_freq.items():
        idf[word] = math.log((N + 1) / (df + 1)) + 1  # 平滑 IDF

    # 为未出现的词提供默认 IDF
    idf["__default__"] = math.log(N + 1) + 1

    _idf_cache = idf
    _idf_cache_chunk_count = total_chunks
    logger.info(f"IDF 缓存已更新，{len(idf)} 个词条")

    return idf


def _format_comparison_context(references: list[dict]) -> str:
    """将参考资料按标准分组，构建对比 prompt 上下文"""
    groups = {}
    for r in references:
        did = r.get("document_id", 0)
        if did not in groups:
            groups[did] = {
                "filename": r.get("filename", "未知标准"),
                "items": [],
            }
        groups[did]["items"].append(r)

    parts = []
    for did, g in groups.items():
        parts.append(f"\n【{g['filename']}】")
        for i, item in enumerate(g["items"], 1):
            section = item.get("section_path", "") or ""
            chunk_type = item.get("chunk_type_cn", "") or item.get("chunk_type", "") or ""
            content = item.get("content", "") or ""
            # 截断过长内容
            if len(content) > 600:
                content = content[:600] + "..."
            parts.append(
                f"  条款{i} [{chunk_type}] {section}\n"
                f"  {content}\n"
            )

    return "\n".join(parts)


def _expand_query_synonyms(question: str) -> str:
    """
    零结果查询扩展：对长查询尝试删除修饰词生成更短的查询。
    不调用 LLM，纯规则。
    """
    # 尝试去掉"是什么""有哪些""如何""怎么"等疑问前缀
    import re
    expanded = question
    # 去掉疑问词
    expanded = re.sub(r'^(请问|请|一下|什么是|什么叫|有哪些|如何|怎么|能否|可以)', '', expanded).strip()
    if expanded and expanded != question:
        return expanded
    return question


def hybrid_search_chunks(db, question: str, limit: int = 5, document_ids: list[int] | None = None):
    """
    混合检索 v2：IDF加权关键词 + 语义相似度 + RRF 融合 + 置信度分层。
    支持多标准检索：document_ids=[1,2,3] 时在多个标准中并行检索。
    返回 (results, confidence) 元组
    """
    terms = extract_search_terms(question)
    idf_weights = _compute_idf_weights(db)

    # ── 关键词检索（IDF 加权 + 多因子评分）──
    keyword_chunks = extract_keyword_terms(db, terms, limit * 3, document_ids=document_ids)

    # 建立关键词排名（用于 RRF）
    kw_ranked = {}  # chunk_id -> rank (1-based)
    kw_scored = {}  # chunk_id -> multi-factor score

    for chunk in keyword_chunks:
        content = chunk.content or ""
        matched_terms = []
        matched_positions = []  # 匹配位置（字符索引）
        total_idf_bonus = 0.0
        exact_phrase_match = False

        for term in terms:
            term_s = term.strip()
            if not term_s:
                continue
            if term_s in content:
                matched_terms.append(term_s)
                # IDF 加权
                idf = idf_weights.get(term_s, idf_weights.get("__default__", 1.0))
                total_idf_bonus += min(idf / 10.0, 0.05)  # 归一化到 0~0.05 每词
                # 记录匹配位置
                pos = content.find(term_s)
                matched_positions.append(pos)

        if not matched_terms:
            continue

        # 完整短语命中（问题原文出现在 chunk 中）
        if question.strip() in content:
            exact_phrase_match = True

        # 匹配密度
        match_density = min(1.0, len(matched_terms) / max(1, len(jieba.lcut(content))))

        # 位置加分（前 200 字符内命中 → 标题/首段命中）
        position_bonus = 0.0
        if matched_positions:
            min_pos = min(matched_positions)
            if min_pos < 200:
                position_bonus = 0.05
            elif min_pos < 500:
                position_bonus = 0.02

        # 多因子评分（0.50 ~ 0.90）
        keyword_score = (
            0.50
            + total_idf_bonus            # IDF 加权 (0 ~ 0.15)
            + match_density * 0.10       # 匹配密度 (0 ~ 0.10)
            + position_bonus             # 位置 (0 ~ 0.05)
            + (0.05 if exact_phrase_match else 0.0)  # 完整短语
        )
        keyword_score = min(0.90, max(0.50, keyword_score))

        kw_scored[chunk.id] = {
            "chunk": chunk,
            "score": keyword_score,
            "match_type": "keyword",
        }

    # 按分数排序给 RRF 排名
    kw_sorted = sorted(kw_scored.values(), key=lambda x: x["score"], reverse=True)
    for rank, item in enumerate(kw_sorted, start=1):
        kw_ranked[item["chunk"].id] = rank

    # ── 语义检索 ──
    semantic_results = semantic_search_chunks(db, question, limit * 3, document_ids=document_ids)
    sem_ranked = {}  # chunk_id -> rank
    sem_scored = {}  # chunk_id -> item

    for rank, item in enumerate(semantic_results, start=1):
        cid = item["chunk"].id
        sem_ranked[cid] = rank
        sem_scored[cid] = item

    # ── RRF 融合 (k=60) ──
    K = 60
    result_map = {}

    all_chunk_ids = set(list(kw_ranked.keys()) + list(sem_ranked.keys()))

    for cid in all_chunk_ids:
        kw_r = kw_ranked.get(cid)
        sem_r = sem_ranked.get(cid)

        rrf_score = 0.0
        if kw_r is not None:
            rrf_score += 1.0 / (K + kw_r)
        if sem_r is not None:
            rrf_score += 1.0 / (K + sem_r)

        # 确定 match_type
        if kw_r is not None and sem_r is not None:
            match_type = "hybrid"
        elif kw_r is not None:
            match_type = "keyword"
        else:
            match_type = "semantic"

        # 取 chunk 数据
        if cid in kw_scored:
            chunk = kw_scored[cid]["chunk"]
        else:
            chunk = sem_scored[cid]["chunk"]

        # 取最高原始分（用于分数展示）
        raw_score = 0.0
        if cid in kw_scored:
            raw_score = max(raw_score, kw_scored[cid]["score"])
        if cid in sem_scored:
            raw_score = max(raw_score, sem_scored[cid]["score"])

        result_map[cid] = {
            "chunk": chunk,
            "score": round(rrf_score * 30, 4),  # RRF normalize to 0~1.0 range
            "raw_score": round(raw_score, 4),
            "match_type": match_type,
        }

    # ── 4维加权置信度判定 ──
    results = list(result_map.values())
    results.sort(key=lambda r: r["score"], reverse=True)
    top_results = results[:limit]

    if not top_results:
        return [], "none", {
            "level": "none", "level_label": "无结果",
            "score_range": [0, 0], "total_matched": 0, "above_threshold": 0,
            "match_breakdown": {}, "intent_match": False,
            "top_section": "", "summary": "未找到任何相关的参考资料"
        }

    hybrid_count = sum(1 for r in top_results if r["match_type"] == "hybrid")
    kw_count = sum(1 for r in top_results if r["match_type"] == "keyword")
    sem_count = sum(1 for r in top_results if r["match_type"] == "semantic")
    top_score = top_results[0]["score"]
    scores = [r["score"] for r in top_results[:3]]
    score_range = [round(min(scores), 2), round(max(scores), 2)]

    # 维度1: 匹配信号强度 (0-1)
    if hybrid_count >= 2:
        signal_score = 1.0
    elif hybrid_count >= 1:
        signal_score = 0.6
    elif sem_count >= 1:
        signal_score = 0.3
    else:
        signal_score = 0.2

    # 维度2: 最高分 (0-1, RRF归一化后范围 0~0.98)
    if top_score >= 0.70:
        top_score_weight = 1.0
    elif top_score >= 0.45:
        top_score_weight = 0.6
    elif top_score >= 0.30:
        top_score_weight = 0.3
    else:
        top_score_weight = 0.0

    # 维度3: 区分度 (0-1)
    score_gap = max(scores) - min(scores) if len(scores) >= 2 else 0
    discrimination = 1.0 if score_gap >= 0.05 else (0.5 if score_gap >= 0.03 else 0.0)

    # 维度4: 意图一致性 (0-1)
    intent = infer_question_intent(question)
    intent_match = False
    top_chunk_type = top_results[0]["chunk"].chunk_type if top_results else None
    intent_type_map = {
        "scope": "scope", "definition": "term", "requirement": "clause",
        "references": "references", "clause_number": "clause",
    }
    expected_type = intent_type_map.get(intent)
    if expected_type and top_chunk_type == expected_type:
        intent_match = True
    intent_consistency = 1.0 if intent_match else 0.3

    # 加权总分
    total_weighted = (
        signal_score * 0.40 +
        top_score_weight * 0.30 +
        discrimination * 0.15 +
        intent_consistency * 0.15
    )

    if total_weighted >= 0.70:
        confidence = "high"
        level_label = "高可信"
    elif total_weighted >= 0.40:
        confidence = "medium"
        level_label = "中等可信"
    else:
        confidence = "low"
        level_label = "低可信，已过滤"

    top_section = (top_results[0]["chunk"].section_path or "") if top_results else ""

    match_breakdown = {}
    if hybrid_count: match_breakdown["hybrid"] = hybrid_count
    if kw_count: match_breakdown["keyword"] = kw_count
    if sem_count: match_breakdown["semantic"] = sem_count

    if confidence == "high":
        summary = f"{hybrid_count}个chunk关键词+语义双命中，最高分{top_score:.2f}，检索与问题意图一致"
    elif confidence == "medium":
        summary = f"部分命中，最高分{top_score:.2f}，建议确认结果相关性"
    else:
        summary = "检索结果可信度低，建议使用更具体的术语或标准编号重新提问"

    confidence_detail = {
        "level": confidence,
        "level_label": level_label,
        "score_range": score_range,
        "total_matched": len(results),
        "above_threshold": len(top_results),
        "match_breakdown": match_breakdown,
        "intent_match": intent_match,
        "top_section": top_section,
        "summary": summary,
    }

    return top_results, confidence, confidence_detail


def expand_search_results(db, matched_chunks: list, depth: int = 1) -> list:
    """
    联动检索：命中 chunk 时，连带返回父子 chunk。
    - depth=1: 返回直接父 + 直接子
    - 父子关系从 chunk_relations 表和 parent_chunk_id 字段获取
    返回扩展后的 chunk 列表（去重）。
    """
    if not matched_chunks:
        return []

    chunk_ids = set()
    result_chunks = []

    for item in matched_chunks:
        chunk = item if isinstance(item, DocumentChunk) else (item.get("chunk") if isinstance(item, dict) else item)
        if chunk and chunk.id not in chunk_ids:
            chunk_ids.add(chunk.id)
            result_chunks.append(chunk)

    if depth <= 0:
        return result_chunks

    # 查询父级和子级
    for chunk in list(result_chunks):
        # 有 parent_chunk_id → 加入父
        if chunk.parent_chunk_id:
            parent = db.query(DocumentChunk).filter(DocumentChunk.id == chunk.parent_chunk_id).first()
            if parent and parent.id not in chunk_ids:
                chunk_ids.add(parent.id)
                result_chunks.append(parent)

        # 查询 chunk_relations 中的子级
        from models import ChunkRelation
        children = db.query(ChunkRelation).filter(
            ChunkRelation.chunk_id == chunk.id,
            ChunkRelation.relation_type == "child"
        ).all()
        for child_rel in children:
            if child_rel.related_chunk_id not in chunk_ids:
                child = db.query(DocumentChunk).filter(DocumentChunk.id == child_rel.related_chunk_id).first()
                if child:
                    chunk_ids.add(child.id)
                    result_chunks.append(child)

    return result_chunks


def normalize_question(question: str) -> str:
    """规范化问题：去除标点和多余空白"""
    if question is None:
        return ""
    question = question.strip()
    question = re.sub(r"[？?！!。，,：:；;、\s]+", "", question)
    return question


def infer_question_intent(question: str) -> str:
    """推断问题意图"""
    q = normalize_question(question)
    # 引用文件意图（放在最前面，优先级最高）
    if any(word in q for word in ["引用", "引用了", "参考了", "参考标准", "规范性引用", "引用的标准", "哪些标准", "引用文件", "引用哪些"]):
        return "references"
    if any(word in q for word in ["适用范围", "适用于", "适合", "适用场景", "哪些场景", "适用对象", "适用哪些"]):
        return "scope"
    if any(word in q for word in ["定义", "什么是", "术语", "含义", "解释", "什么叫", "是指", "的概念"]):
        return "definition"
    if any(word in q for word in ["要求", "必须", "应当", "不得", "指标", "参数", "规定", "标准值", "规范", "阈值", "检测", "检验"]):
        return "requirement"
    # 条款编号意图（问题中包含"第X条"/"X.X"/"第X章"等）
    if re.search(r'第[一二三四五六七八九十百\d]+[条款章]', q) or re.search(r'\d+\.\d+', q):
        return "clause_number"
    return "general"

# 用于从 chunk 内容中去掉内置前缀（标准/路径/页码已在外部显示）
_CONTENT_PREFIX_RE = re.compile(
    r'^(?:标准[：:][^\n]*\n)?(?:章节路径[：:][^\n]*\n)?(?:页码[：:][^\n]*\n)?---\n?',
    re.MULTILINE
)

def _strip_chunk_prefix(content: str) -> str:
    """去掉 chunk 内置的前缀行，避免在 context 中重复显示"""
    return _CONTENT_PREFIX_RE.sub('', content).strip()


# ═══════════════════════════════════════════════════════════════
# 相关标准推荐（多维度融合）
# ═══════════════════════════════════════════════════════════════

def recommend_related_standards(
    db,
    current_document_id: int,
    question: str,
    limit: int = 5
) -> dict:
    """
    基于当前文档和用户问题，推荐相关标准。
    三种推荐策略：
    1. 同行业推荐 — 相同 industry 标签的其他标准
    2. 语义相似推荐 — 问题与各标准前言/范围 chunk 的向量相似度
    3. 显式引用推荐 — 当前标准"规范性引用文件"中引用的标准编号
    """
    current_doc = get_document_by_id(db, current_document_id)
    if current_doc is None:
        return {"same_industry": [], "semantic_similar": [], "cited_references": []}

    recommendations = {
        "same_industry": [],
        "semantic_similar": [],
        "cited_references": [],
    }

    # ── 策略1：同行业推荐 ──
    if current_doc.industry:
        same_industry_docs = db.query(Document).filter(
            Document.industry == current_doc.industry,
            Document.id != current_document_id
        ).limit(limit).all()

        recommendations["same_industry"] = [
            {
                "document_id": doc.id,
                "filename": doc.filename,
                "standard_type": doc.standard_type,
                "industry": doc.industry,
                "reason": "同属「{}」行业标准".format(current_doc.industry),
            }
            for doc in same_industry_docs
        ]

    # ── 策略2：语义相似推荐 ──
    try:
        question_emb = services.embedding_service.generate_embedding(question)

        # 扫描所有其他文档的范围/前言/cov chunk
        other_chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id != current_document_id,
            DocumentChunk.chunk_type.in_(["scope", "preface", "cover"]),
            DocumentChunk.embedding.isnot(None),
        ).all()

        scored_docs = {}
        for chunk in other_chunks:
            chunk_emb = services.embedding_service.json_to_embedding(chunk.embedding)
            score = services.embedding_service.cosine_similarity(question_emb, chunk_emb)
            if score > 0.55:  # 阈值
                did = chunk.document_id
                if did not in scored_docs or scored_docs[did]["score"] < score:
                    doc = get_document_by_id(db, did)
                    if doc:
                        scored_docs[did] = {
                            "document_id": did,
                            "filename": doc.filename,
                            "standard_type": doc.standard_type,
                            "industry": doc.industry,
                            "score": round(score, 3),
                            "reason": "内容与您的问题语义相似（匹配度 {:.0%}）".format(score),
                        }

        recommendations["semantic_similar"] = sorted(
            scored_docs.values(), key=lambda x: x["score"], reverse=True
        )[:limit]
    except Exception:
        pass  # embedding 计算失败时降级，不阻塞

    # ── 策略3：规范性引用文件推荐 ──
    ref_chunks = db.query(DocumentChunk).filter(
        DocumentChunk.document_id == current_document_id,
        DocumentChunk.chunk_type == "references",
    ).all()

    cited_numbers = set()
    for ref_chunk in ref_chunks:
        raw_content = ref_chunk.content or ""
        clean = _strip_chunk_prefix(raw_content)
        # 复用已有的标准编号正则
        for match in _STD_NUM_RE.finditer(clean):
            cited_numbers.add(match.group(0).strip())

    for cited_num in cited_numbers:
        # 在库内按文件名模糊匹配
        matched_docs = db.query(Document).filter(
            Document.filename.contains(cited_num)
        ).all()
        for doc in matched_docs:
            if doc.id != current_document_id:
                recommendations["cited_references"].append({
                    "document_id": doc.id,
                    "filename": doc.filename,
                    "standard_type": doc.standard_type,
                    "industry": doc.industry,
                    "cited_number": cited_num,
                    "reason": "当前标准「规范性引用文件」中引用了 {}".format(cited_num),
                })

    # 去重（同文档可能被多策略命中）
    seen = set()
    for key in ["same_industry", "semantic_similar", "cited_references"]:
        unique = []
        for item in recommendations[key]:
            if item["document_id"] not in seen:
                seen.add(item["document_id"])
                unique.append(item)
        recommendations[key] = unique

    return recommendations


# ═══════════════════════════════════════════════════════════════
# Grounded Evidence Extraction — 三层防御（防幻觉）
# ═══════════════════════════════════════════════════════════════

def _split_sentences_zh(text: str) -> list[str]:
    """拆分中文句子，保护编号中的小数点（如 3.3.1 不会在中间断开）"""
    # 临时保护数字中的小数点
    text = re.sub(r'(\d)\.(\d)', r'\1__DOT__\2', text)
    # 按中文标点切分（。！？；）
    parts = re.split(r'(?<=[。！？；])\s*', text)
    result = []
    for part in parts:
        part = part.replace('__DOT__', '.').strip()
        if part and len(part) >= 8:
            result.append(part)
    return result


def extract_evidence_sentences(question: str, references: list[dict],
                                similarity_threshold: float = 0.45) -> list[dict]:
    """
    层1：从检索到的 reference chunks 中抽取与问题最相关的证据句子（非LLM）。

    流程：
    1. 将每个 reference 的 content 按句子拆分
    2. 批量计算所有句子与 question 的余弦相似度
    3. 保留相似度 > threshold 的句子
    4. 每句标注来源元数据

    返回: [{"sentence_id": int, "text": str, "similarity": float,
            "chunk_index": int, "document_id": int, "section_number": str,
            "section_path": str, "page_number": int, "chunk_type": str,
            "filename": str}, ...]
    """
    if not references or not question:
        return []

    # 1. 拆分所有句子
    all_sentences = []
    for ri, ref in enumerate(references):
        content = ref.get("content", "")
        if not content:
            continue
        sentences = _split_sentences_zh(content)
        for si, sent in enumerate(sentences):
            all_sentences.append({
                "text": sent,
                "ref_index": ri,
                "local_index": si,
                "document_id": ref.get("document_id"),
                "chunk_index": ref.get("chunk_index"),
                "section_number": ref.get("section_number"),
                "section_path": ref.get("section_path"),
                "page_number": ref.get("page_number"),
                "chunk_type": ref.get("chunk_type"),
                "filename": ref.get("filename"),
            })

    if not all_sentences:
        return []

    # 2. 批量计算相似度
    from services import embedding_service
    q_embedding = embedding_service.generate_query_embedding(question)
    sent_texts = [s["text"] for s in all_sentences]
    sent_embeddings = embedding_service.generate_embeddings_batch(sent_texts)

    # 3. 筛选（带降级：主阈值 0.45 无结果时降到 0.35）
    for attempt_threshold in (similarity_threshold, similarity_threshold - 0.10):
        results = []
        for idx, (emb, sent) in enumerate(zip(sent_embeddings, all_sentences)):
            sim = embedding_service.cosine_similarity(q_embedding, emb)
            if sim >= attempt_threshold:
                results.append({
                    "sentence_id": idx,
                    "text": sent["text"],
                    "similarity": round(sim, 4),
                    "document_id": sent["document_id"],
                    "chunk_index": sent["chunk_index"],
                    "section_number": sent["section_number"],
                    "section_path": sent["section_path"],
                    "page_number": sent["page_number"],
                    "chunk_type": sent["chunk_type"],
                    "filename": sent["filename"],
                })
        if results:
            results.sort(key=lambda x: x["similarity"], reverse=True)
            return results[:30]
    return []


def assemble_verified_answer(selection: dict, sentences: list[dict]) -> dict:
    """
    层3：根据 LLM 选择的句子组装最终答案，逐句验证可追溯到原文。

    参数:
      selection: LLM 返回的选择结果 {"selected": [...], "not_found": [...], "summary": "..."}
      sentences: 层1 抽取的证据句子列表

    返回: {"answer": str, "verified": bool, "citations": [...], "not_found": [...],
            "evidence_count": int, "selected_count": int}
    """
    sent_map = {s["sentence_id"]: s for s in sentences}
    selected = selection.get("selected", [])
    not_found = selection.get("not_found", [])
    summary = selection.get("summary", "")

    if not selected and not summary:
        return {
            "answer": "当前资料中未找到与该问题直接相关的内容。建议：1) 尝试用更具体的术语重新提问 2) 确认相关标准文件已上传。",
            "verified": True,
            "citations": [],
            "not_found": not_found,
            "evidence_count": len(sentences),
            "selected_count": 0,
        }

    # 组装
    parts = []
    citations = []

    for item in selected:
        sid = item.get("sentence_id")
        if sid is None:
            continue
        sent = sent_map.get(sid)
        if sent is None:
            continue

        text = sent["text"]
        relevance = item.get("relevance", "")

        # 构建来源标记
        source_parts = []
        if sent.get("section_number"):
            source_parts.append("第{}条".format(sent["section_number"]))
        if sent.get("section_path"):
            source_parts.append(sent["section_path"])
        if sent.get("page_number"):
            source_parts.append("第{}页".format(sent["page_number"]))
        source = "「{}」".format(" | ".join(source_parts)) if source_parts else ""

        parts.append({
            "text": text,
            "source": source,
            "relevance": relevance,
        })
        citations.append({
            "sentence_id": sid,
            "text": text,
            "section_path": sent.get("section_path"),
            "section_number": sent.get("section_number"),
            "page_number": sent.get("page_number"),
            "document_id": sent.get("document_id"),
            "chunk_index": sent.get("chunk_index"),
            "filename": sent.get("filename"),
            "similarity": sent.get("similarity"),
            "relevance": relevance,
        })

    # 构建最终答案
    answer = ""
    if summary:
        answer += "**📋 检索摘要：**{}\n\n---\n\n".format(summary)

    for i, part in enumerate(parts):
        answer += "**{}** {} {}\n\n".format(i + 1, part["text"], part["source"])

    if not_found:
        answer += "---\n**⚠️ 当前资料中未找到以下信息：**\n"
        for nf in not_found:
            answer += "- {}\n".format(nf)
        answer += "\n建议核实原文或上传更完整的标准文件。"

    return {
        "answer": answer,
        "verified": True,
        "citations": citations,
        "not_found": not_found,
        "evidence_count": len(sentences),
        "selected_count": len(parts),
    }
